import getpass
import os
from langchain_nvidia_ai_endpoints import ChatNVIDIA
from langchain_nvidia_ai_endpoints import NVIDIAEmbeddings
from docx import Document
# Here we create a vector store from the documents and save it to disk.
from operator import itemgetter
from langchain.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain.text_splitter import CharacterTextSplitter
from langchain_nvidia_ai_endpoints import ChatNVIDIA
import faiss
import gradio as gr

NIMKEYConfig = 'D:\CodeRepositories\AILearning\NIMKey'
if os.environ.get("NVIDIA_API_KEY", "").startswith("nvapi-"):
    print("Valid NVIDIA_API_KEY already in environment. Delete to reset")
else:
    nvapi_key = None
    with open(NIMKEYConfig, 'r') as file:
        nvapi_key = file.read()
    
    assert nvapi_key.startswith("nvapi-"), f"{nvapi_key[:5]}... is not a valid key"
    os.environ["NVIDIA_API_KEY"] = nvapi_key


llm = ChatNVIDIA(model="microsoft/phi-3-vision-128k-instruct", nvidia_api_key=nvapi_key, max_tokens=512)
embedder = NVIDIAEmbeddings(model="NV-Embed-QA")

def LoadDocument():
    dataPath = "./data/"
    ps = os.listdir(dataPath)
    data = []
    sources = []
    for p in ps:
        if p.endswith(".docx"):
            doc = Document(dataPath + p)
            for para in doc.paragraphs:
                if len(para.text) > 0:
                    data.append(para.text)
                    sources.append(p)

    print(data)
    print(sources)

    documents=[d for d in data if d != '\r\n']
    # 只需要执行一次，后面可以重读已经保存的向量存储
    text_splitter = CharacterTextSplitter(chunk_size=400, separator=" ")
    docs = []
    metadatas = []
    for i, d in enumerate(documents):
        splits = text_splitter.split_text(d)
        #print(len(splits))
        docs.extend(splits)
        metadatas.extend([{"source": sources[i]}] * len(splits))
    store = FAISS.from_texts(docs, embedder , metadatas=metadatas)
    store.save_local(f'{dataPath}nv_embedding')

def get_response(message, history):# Load the vectorestore back.
    dataPath = "./data/"
    store = FAISS.load_local(f"{dataPath}nv_embedding", embedder,allow_dangerous_deserialization=True)
    retriever = store.as_retriever()

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Answer solely based on the following context:\n<Documents>\n{context}\n</Documents>",
            ),
            ("user", "{question}"),
        ]
    )

    chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    return chain.invoke(message)
    # return random.choice(["好的，我知道了", "谢谢你的提醒", "我会注意的", "谢谢你的建议"])

gr.ChatInterface(get_response).launch()